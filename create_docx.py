from pathlib import Path
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('/mnt/data/LeadFlowML_Project')
OUT = ROOT / 'Executive_Summary_LeadFlowML.docx'

report_default = json.loads((ROOT / 'outputs' / 'evaluation_report.json').read_text())
report_rf = json.loads((ROOT / 'outputs' / 'evaluation_report_rf.json').read_text())


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=10.5, color='000000'):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles['Normal'].font.name = 'Aptos'
styles['Normal'].font.size = Pt(10.5)
styles['Title'].font.name = 'Aptos Display'
styles['Title'].font.size = Pt(24)
styles['Heading 1'].font.name = 'Aptos Display'
styles['Heading 1'].font.size = Pt(15)
styles['Heading 2'].font.name = 'Aptos Display'
styles['Heading 2'].font.size = Pt(12)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('LeadFlowML DSL\n')
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x17, 0x2B, 0x4D)
r2 = p.add_run('Executive Summary of Design and Implementation')
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x44, 0x55, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Domain: Sales lead qualification with machine-learning case classification')
r.italic = True
r.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Deliverable context: textX DSL + CWL workflow + notebook demo + presentation package')
r.font.size = Pt(10.5)

# Executive summary block
summary = doc.add_table(rows=1, cols=1)
summary.alignment = WD_TABLE_ALIGNMENT.CENTER
summary.autofit = True
cell = summary.cell(0,0)
shade_cell(cell, 'EAF2FF')
set_cell_text(cell,
    'LeadFlowML is a lightweight domain-specific language built with textX for defining an end-to-end machine-learning workflow in a sales scenario. '
    'A short DSL specification describes the training dataset, new-case dataset, feature groups, classification model, threshold, metrics, and output artifacts. '
    'The processor parses this DSL into a Python object model and generates CWL CommandLineTools plus a CWL Workflow, allowing the same business specification to run locally or through a portable workflow engine.',
    size=11)

doc.add_paragraph('')

# Section 1
h = doc.add_paragraph(style='Heading 1')
h.add_run('1. Problem Framing')
doc.add_paragraph(
    'The assignment requires a DSL that is specialized for a single problem area and abstracts away general programming noise. '
    'To make the domain concrete, this project uses a sales lead qualification scenario where historical CRM cases act as training data and incoming leads are classified as qualified or not qualified. '
    'The target user is an ML expert or analyst who wants to express the workflow declaratively rather than wiring together scripts by hand.'
)
doc.add_paragraph(
    'The design also follows the workflow ideas emphasized in the CWL and workflow-engine literature: clear tool boundaries, explicit inputs/outputs, reuse of command-line components, and reproducibility through structured workflow files.'
)

# Section 2 design goals
h = doc.add_paragraph(style='Heading 1')
h.add_run('2. DSL Design Goals')
items = [
    'Readable by domain experts: the syntax resembles a compact project specification rather than general Python code.',
    'Traceable execution: every important business choice (algorithm, threshold, features, file paths, metrics) is visible in the DSL source.',
    'Portable orchestration: the DSL is converted into CWL artifacts so the workflow can be validated and executed in a standard way.',
    'Extensible structure: the current grammar can later be extended with hyperparameter sweeps, model selection, or deployment targets.'
]
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

# Section 3 syntax table
h = doc.add_paragraph(style='Heading 1')
h.add_run('3. Syntax and Semantics')
doc.add_paragraph(
    'LeadFlowML uses a textual grammar with the following top-level sections. Each section maps directly to an execution concern in the generated workflow.'
)

tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
headers = ['DSL section', 'Purpose', 'Generated effect']
for i, text in enumerate(headers):
    shade_cell(tbl.cell(0,i), 'D9E7FF')
    set_cell_text(tbl.cell(0,i), text, bold=True, size=10.5, color='17324D')
rows = [
    ('data', 'Identifies training and scoring CSV files plus id/target columns', 'Becomes workflow inputs and validation rules'),
    ('features', 'Separates numeric and categorical variables', 'Used by preprocessing and model pipeline generation'),
    ('model', 'Defines algorithm, test split, random state, threshold, CV folds', 'Controls training/evaluation behavior and output scoring logic'),
    ('metrics', 'Lists evaluation measures expected by the analyst', 'Included in the generated evaluation report'),
    ('outputs', 'Declares the desired artifact locations', 'Controls model, prediction, and report file names in the generated CWL tools'),
]
for row in rows:
    cells = tbl.add_row().cells
    for i, text in enumerate(row):
        set_cell_text(cells[i], text)

# code snippet
h = doc.add_paragraph(style='Heading 2')
h.add_run('Example DSL snippet')
code_lines = [
'project LeadQualification',
'description "Classify new sales leads from historical CRM cases using a reproducible workflow."',
'',
'data',
'  train "data/train_sales_cases.csv"',
'  score "data/new_sales_cases.csv"',
'  id lead_id',
'  target qualified',
'',
'features',
'  numeric age, annual_income, website_visits, prior_purchases, support_tickets',
'  categorical region, channel, loyalty_segment',
'',
'model',
'  algorithm logistic_regression',
'  test_size 0.25',
'  random_state 42',
'  threshold 0.55',
'  cv_folds 5',
]
for line in code_lines:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(line)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)

# Section 4 Implementation
h = doc.add_paragraph(style='Heading 1')
h.add_run('4. Implementation Architecture')
arch = [
    ('textX grammar', 'Defines the external DSL and builds the Python model.'),
    ('DSL processor', 'Parses the specification and emits CWL YAML files plus a JSON summary.'),
    ('ML runtime scripts', 'Validate data, train/evaluate the model, and classify new cases.'),
    ('CWL workflow', 'Connects validation, training, and scoring as reusable command-line tools.'),
    ('Notebook demo', 'Shows multiple usage scenarios and compares logistic regression vs random forest.'),
]
for name, desc in arch:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{name}: ').bold = True
    p.add_run(desc)

doc.add_paragraph(
    'The implementation produced a working end-to-end prototype. The default scenario trains a logistic-regression classifier; an alternate scenario swaps in a random-forest model without changing any Python source code, demonstrating the value of declarative configuration.'
)

# Section 5 Results
h = doc.add_paragraph(style='Heading 1')
h.add_run('5. Demonstrated Results')
metrics_table = doc.add_table(rows=1, cols=3)
metrics_table.style = 'Table Grid'
metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, text in enumerate(['Metric', 'Logistic regression', 'Random forest']):
    shade_cell(metrics_table.cell(0,i), 'D9E7FF')
    set_cell_text(metrics_table.cell(0,i), text, bold=True, color='17324D')
metrics = ['accuracy', 'precision', 'recall', 'f1', 'roc_auc', 'cv_f1_mean']
for m in metrics:
    row = metrics_table.add_row().cells
    set_cell_text(row[0], m)
    set_cell_text(row[1], f"{report_default['metrics'][m]:.3f}")
    set_cell_text(row[2], f"{report_rf['metrics'][m]:.3f}")

img = ROOT / 'artifacts' / 'metrics_comparison.png'
if img.exists():
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img), width=Inches(6.0))
    p2 = doc.add_paragraph('Figure: Metric comparison generated from the demo notebook.')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].italic = True
    p2.runs[0].font.size = Pt(9)

# Section 6 alignment with literature
h = doc.add_paragraph(style='Heading 1')
h.add_run('6. Alignment with the Referenced Literature')
doc.add_paragraph(
    'The design reflects the textX approach of defining the grammar and resulting model structure from a common meta-language, which keeps the implementation compact and Python-native. '
    'It also follows the CWL model of separating reusable CommandLineTools from Workflow orchestration, making inputs, outputs, and step dependencies explicit. '
    'Finally, the workflow inspiration from the scientific-pipeline literature is visible in the reusable component boundaries, the focus on reproducibility, and the support for alternative ML scenarios without rewriting the orchestration logic.'
)

# Section 7 deliverables
h = doc.add_paragraph(style='Heading 1')
h.add_run('7. Package Contents')
contents = [
    'Executive summary in MS Word',
    'Executed Jupyter notebook with scenario demonstrations',
    'Full project package with source files, grammar, workflows, generated CWL, data, outputs, and README',
    'PowerPoint presentation for the group presentation'
]
for c in contents:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(c)

# References
h = doc.add_paragraph(style='Heading 1')
h.add_run('References')
refs = [
    'Tomassetti, F. Quick Domain-Specific Languages in Python with textX. https://tomassetti.me/domain-specific-languages-in-python-with-textx/',
    'Karle, N. et al. Parsl+CWL: Towards Combining the Python and CWL Ecosystems. https://arxiv.org/html/2412.08062v1',
    'Lampa, S. et al. SciPipe: A workflow library for agile development of complex and dynamic bioinformatics pipelines. https://www.diva-portal.org/smash/get/diva2:1242254/FULLTEXT02.pdf'
]
for r in refs:
    p = doc.add_paragraph(style='List Number')
    p.add_run(r)

doc.save(OUT)
print(OUT)
